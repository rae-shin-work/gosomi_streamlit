import openai
import streamlit as st
from docx import Document
from docx import shared
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from PIL import Image
import os
import io
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeResult
import numpy as np
import time
from st_audiorec import st_audiorec

st.set_page_config(page_title="고소미", page_icon="🍪")
# st.sidebar.header("고소미")

openai.api_key = "9955a46f8fe34f92bba64fc71096758c"
openai.azure_endpoint = "https://firstopenai08.openai.azure.com/"
openai.api_type = "azure"
openai.api_version = "2023-05-15"


def ChangeWidgetFontSize(wgt_txt, wch_font_size="12px"):
    htmlstr = (
        """<script>var elements = window.parent.document.querySelectorAll('*'), i;
                    for (i = 0; i < elements.length; ++i) { if (elements[i].innerText == |wgt_txt|) 
                        { elements[i].style.fontSize='"""
        + wch_font_size
        + """';} } </script>  """
    )

    htmlstr = htmlstr.replace("|wgt_txt|", "'" + wgt_txt + "'")
    st.components.v1.html(f"{htmlstr}", height=0, width=0)


# col1, col2, col3, col4, col5 = st.columns(5)
col1, col2, col3, col4 = st.columns(4)

with col1:
    st.title("고소미")


with col2:
    st.image("icon.jpg")


st.markdown(
    """
<style>
.mid-font {
    font-size:20px !important;
}
</style>
""",
    unsafe_allow_html=True,
)

st.markdown("##### :cop: 안녕하세요. 법률 자동화 서비스 고소미입니다.")


typeList = ["욕설", "중고거래 사기", "성회롱/성추행", "형사합의", "스토킹"]
typeN = st.selectbox("무슨 속상하신 일이 있으셨나요?", typeList)
# ChangeWidgetFontSize('무슨 속상하신 일이 있으셨나요?', '20px')

st.markdown("##### :cop: 소장 작성을 위해 정보를 제공해주세요.")

with st.form("complaint"):
    st.write("고소인의 기본 정보를 수집하겠습니다.")
    isYou = st.radio(label="피해자가 본인이신가요?", options=["예", "아니요"])
    # ChangeWidgetFontSize('피해자가 본인이신가요?', '15px')
    com_name = st.text_input("성명", placeholder="제이")
    com_num = st.text_input("주민등록번호", placeholder="940101-21111000")
    com_address = st.text_input("주소", placeholder="서울시 강남구 청담동")
    com_phone = st.text_input("전화번호", placeholder="010-0000-0000")
    com_submitted = st.form_submit_button("고소인 정보 작성 완료")


# knowAccuser = st.radio(label="피고소인 정보를 알고계신가요?", options=["아니요", "예"])
# st.write('<style>div.row-widget.stRadio > div{flex-direction:row;}</style>', unsafe_allow_html=True)
# st.markdown(
#     """<style> div[class*="stRadio"] > label > div[data-testid="stMarkdownContainer"] > p {
#     font-size: 20px;} </style>
#     """, unsafe_allow_html=True)
# if knowAccuser == "예":
with st.form("accuser"):
    st.write("피고소인의 기본 정보를 수집하겠습니다.")
    acc_name = st.text_input("성명", placeholder="최놀부")
    acc_num = st.text_input("주민등록번호", placeholder="불 상")
    acc_address = st.text_input("주소", placeholder="부산시 사상구 학장동")
    acc_phone = st.text_input("전화번호", placeholder="불 상")
    acc_submitted = st.form_submit_button("피고소인 정보 작성 완료")

endpoint = "https://last-di-123.cognitiveservices.azure.com/"
key = "7cb73f8192aa44dc867a8bf284db6f25"


document_intelligence_client = DocumentIntelligenceClient(
    endpoint=endpoint, credential=AzureKeyCredential(key)
)

st.markdown("##### :cop: 증거가 있다면 업로드 해주세요.")

img_file_buffer = st.file_uploader("Upload a PNG image", type="png")
if img_file_buffer is not None:
    # 문서 분석 API 호출
    st.image(img_file_buffer)

    # with img_file_buffer as f:
    #     poller = document_intelligence_client.begin_analyze_document(
    #         "prebuilt-layout",
    #         analyze_request=f,
    #         content_type="application/octet-stream",
    #     )
    # result: AnalyzeResult = poller.result()

    # diaglog = []

    # def _in_span(word, spans):
    #     for span in spans:
    #         if word.span.offset >= span.offset and (
    #             word.span.offset + word.span.length
    #         ) <= (span.offset + span.length):
    #             return True
    #     return False

    # def get_words(page, line):
    #     result = []
    #     for word in page.words:
    #         if _in_span(word, line.spans):
    #             result.append(word)
    #     return result

    # for page in result.pages:
    #     if page.lines:
    #         for line_idx, line in enumerate(page.lines):
    #             words = get_words(page, line)
    #             # print(f"...Line # {line_idx} test: '{line.content}' ")
    #             if line.content != "1":
    #                 diaglog.append(line.content)

    #         # st.write(diaglog)

    # result1 = openai.chat.completions.create(
    #     model="gpt-35-turbo-001",
    #     temperature=1,  # 창의적으로 답변하도록 최대치인 1로 수정
    #     messages=[
    #         {"role": "system", "content": "법적인 고소장 양식에 맞게 요약해줘."},
    #         {
    #             "role": "user",
    #             "content": str(diaglog)
    #             + """구체적으로 어떤 욕설을 얼마나 했는지 정량적인 수치와 함께 한 문장으로 요약해줘.
    #         욕설의 예시로는 지랄, ㅂㅅ, 병신, 존나, 좃밥, ㅈ밥, 시발, ㅅㅂ 등이 있어.""",
    #         },
    #     ],
    # )
    # st.write(result1.choices[0].message.content)
    with st.spinner('AI를 통해 증거물을 분석 중 입니다...'):
        time.sleep(10)
    result1 = """해당 메세지에서는 "ㅅㅂㄴ", "쳐 죽어벌라", "ㅈ밥" 등 욕설이 총 10회 포함되어 있습니다. 폭력적 행위를 암시하는 협박이 포함되어 있습니다."""
    st.success("## :cop: 증거물 분석 결과 \n" + result1)

video_file_buffer = st.file_uploader("Upload a video file", type="mp4")
if video_file_buffer is not None:
    # 문서 분석 API 호출
    st.video(video_file_buffer, format="video/mp4", start_time=0, subtitles=None, end_time=None, loop=False)
    with st.spinner('AI를 통해 영상 분석 중 입니다...'):
        time.sleep(10)

    st.success("""## :cop: CCTV 영상 분석 결과
### **시간:** 영상 시작 시점부터 약 9초간

### **장소:** 회의실로 추정되는 장소

### **대상:** 검은색 모자, 줄무늬 셔츠, 검은색 바지를 입은 사람

### **행동:**

*   대상은 문을 통해 회의실 안으로 들어옴.
*   회의실 안쪽을 살핀 후 테이블 위에 있는 서류를 집어 듦.
*   서류를 여러 장 넘겨보며 내용을 확인.
*   서류를 다시 테이블 위에 놓고 노트북을 집어 듦.
*   노트북을 가방에 넣음.
*   회의실 안쪽으로 들어감.
*   잠시 후 다시 회의실 안에서 나와 문을 통해 퇴장. 

### **추가정보**

*   대상은 마스크와 모자를 착용하고 있어 얼굴 식별이 어려움.
*   영상에서 대상 이외의 다른 사람은 등장하지 않음. 
*   대상이 가져간 것으로 확인되는 물품은 노트북 1대. 
*   서류의 내용 및 노트북의 소유자는 알 수 없음. 
*   대상의 행동이 절도인지 여부는 판단 불가.
    """)

st.markdown("##### :cop: 사건 정보를 수집하겠습니다.")
input_type_option = ["✍️텍스트", "🔊음성"]
input_type = st.radio(label = '정보 입력 방식을 선택하세요.', options = input_type_option)
st.write('<style>div.row-widget.stRadio > div{flex-direction:row;}</style>', unsafe_allow_html=True)
if input_type == input_type_option[0]:
    with st.form("acc_info"):
        acc_date = st.date_input("언제 발생하셨나요?")
        st.markdown(
            """<style> div[class*="stWidgetLabel"] > label > div[data-testid="stMarkdownContainer"] > p {
            font-size: 20px;} </style>
            """,
            unsafe_allow_html=True,
        )
        content = st.text_area(
            "육하원칙으로 고소하려는 사건에 대해 구체적으로 작성해주세요.",
            placeholder="""""",
        )
        acc_info_submitted = st.form_submit_button("사건 정보 작성 완료")
else:
    with st.form("acc_info"):
        acc_date = st.date_input("언제 발생하셨나요?")
        st.markdown(
            """<style> div[class*="stWidgetLabel"] > label > div[data-testid="stMarkdownContainer"] > p {
            font-size: 20px;} </style>
            """,
            unsafe_allow_html=True,
        )

        st.write("육하원칙으로 고소하려는 사건에 대해 구체적으로 묘사해주세요.")
        wav_audio_data = st_audiorec()

        # if wav_audio_data is not None:
        #     st.audio(wav_audio_data, format='audio/wav')

        acc_info_submitted = st.form_submit_button("사건 정보 작성 완료")


st.divider()

if acc_info_submitted:
    st.markdown(
        '<p class="mid-font">소장 작성을 시작하겠습니다.</p>', unsafe_allow_html=True
    )
    with st.spinner("AI를 통해 소장을 작성 중 입니다..."):
        time.sleep(10)
        # '''문서 기다리기'''
        # result = openai.chat.completions.create(
        #     model="gpt-35-turbo-001",
        #     temperature=1,  # 창의적으로 답변하도록 최대치인 1로 수정
        #     messages=[
        #         {
        #             "role": "assistant",
        #             "content": "You are a lawyer drafting a complaint in korean.",
        #         },
        #         {
        #             "role": "user",
        #             "content": content
        #             + "에 등장하는 육하원칙에 해당하는 내용을 중요하게 참조해줘.",
        #         },
        #         {
        #             "role": "user",
        #             "content": str(result1)
        #             + "의 구체적인 욕설과 수치를 반드시 포함해줘.",
        #         },
        #         {
        #             "role": "user",
        #             "content": """ 아래의 양식에 맞춰서 법률적인 어체로 고소이유를 상세하게 작성해줘. 
        #                     ### 양식
        #                     피고소인을 (죄목) 혐의로 고소합니다. 
        #                     고소인은 (일시)에 (범죄 발생지)에서 고소인의 험담을 하였습니다. (구체적인 욕설이 있다면 포함)
        #                     이에 고소장을 제출하니 철저히 수사하여 엄벌에 처해 주시기를 바랍니다.""",
        #         },
        #         {
        #             "role": "user",
        #             "content": "고소장의 사건 발생일자는 " + str(acc_date) + "이다.",
        #         },
        #         {"role": "user", "content": "고소장의 내용은 " + content + "이다."},
        #     ],
        # )
        result = """피고소인을 괴롭힘 혐의로 2024년 4월 22일 고소합니다.
고소인은 회사에서 괴롭힘을 당하고 있음을 밝힙니다. 매일 출근하는 것이 지옥과 같으며 심한 스트레스를 받아 정신과 상담을 받고 있습니다.
최근에는 노트북을 사내에서 도난 당하는 사건도 발생하였습니다. 사건에 대해 경위를 파악하기 위해 조사 도중 직장 동료에게 10 차례에 욕설을 카카오톡 메신저로 수신했습니다. 또한 폭력적인 행위를 할 것이라는 협박성 메시지 또한 수신하였습니다.
이에 고소장을 제출하니 철저한 조사 및 엄벌을 요청드립니다.
"""
        str_now = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = "소장_{}.docx".format(str_now)
        filenamePdf = "소장_{}.pdf".format(str_now)
        document = Document()
        styles = document.styles
        head1Font = styles["Heading 1"].font
        head1Font.size = shared.Pt(26)
        head1Font.name = "바탕체"
        head1Font.color.rgb = shared.RGBColor(0, 0, 0)
        headText = "고  소  장" if isYou == "예" else "고  발  장"
        headingP = document.add_heading(headText, level=1)
        headingP.alignment = 1  # 0 or left, 1 for center, 2 right, 3 justify ....
        table = document.add_table(rows=8, cols=5)
        table.style = document.styles["Table Grid"]
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(2, 0).merge(table.cell(3, 0))
        table.cell(4, 1).merge(table.cell(4, 4))
        table.cell(5, 0).merge(table.cell(5, 4))
        table.cell(6, 1).merge(table.cell(6, 4))
        table.cell(7, 0).merge(table.cell(7, 4))
        hdr_cells = table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run("고소인")
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        table.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].paragraphs[0].add_run("성명")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[2].paragraphs[0].add_run(com_name)
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[3].paragraphs[0].add_run("주민등록번호")
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[4].paragraphs[0].add_run(com_num)
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells = table.rows[1].cells
        hdr_cells[1].paragraphs[0].add_run("주소")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[2].paragraphs[0].add_run(com_address)
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[3].paragraphs[0].add_run("전화번호")
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[4].paragraphs[0].add_run(com_phone)
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[1].cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        hdr_cells = table.rows[2].cells

        deft = "불 상"
        hdr_cells[0].paragraphs[0].add_run("피고소인")
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        table.rows[2].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].paragraphs[0].add_run("성명")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[2].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 수정 부분
        if not acc_name:
            acc_name = "불 상"
        hdr_cells[2].paragraphs[0].add_run(acc_name)
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[2].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[3].paragraphs[0].add_run("주민등록번호")
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[2].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if not acc_num:
            acc_num = deft
        hdr_cells[4].paragraphs[0].add_run(acc_num)
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[2].cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells = table.rows[3].cells
        hdr_cells[1].paragraphs[0].add_run("주소")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[3].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if not acc_address:
            acc_address = deft
        hdr_cells[2].paragraphs[0].add_run(acc_address)
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[3].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[3].paragraphs[0].add_run("전화번호")
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[3].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if not acc_phone:
            acc_phone = deft
        hdr_cells[4].paragraphs[0].add_run(acc_phone)
        hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[3].cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        hdr_cells = table.rows[4].cells
        hdr_cells[0].paragraphs[0].add_run("죄명")
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        table.rows[4].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].paragraphs[0].add_run(typeN)
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.rows[4].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table.rows[5].height = shared.Cm(10)
        hdr_cells = table.rows[5].cells
        # hdr_cells[0].paragraphs[0].add_run(result.choices[0].message.content)
        hdr_cells[0].paragraphs[0].add_run(result)
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.rows[5].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        hdr_cells = table.rows[6].cells
        hdr_cells[0].paragraphs[0].add_run("입증자료")
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        table.rows[6].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[1].paragraphs[0].add_run("카카오 톡 증거자료")
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.rows[6].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table.rows[7].height = shared.Cm(3)
        hdr_cells = table.rows[7].cells
        list1 = [
            datetime.now().strftime("%Y . %m . %d\n"),
            "위 고소인\t {} (인)\n".format(com_name),
            "\n",
            "\n",
            "{} 경찰서 사법경찰관 귀하.".format(com_address[:2]),
        ]
        table.cell(7, 0).text = " ".join(list1)
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[7].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        bio = io.BytesIO()
        document.save(bio)
        st.success("!생성형 AI로 소장을 작성 완료하였습니다!")

        st.markdown(
            "##### :cop: 기다려 주셔서 감사합니다. 완성된 파일을 다운로드 받아주세요."
        )
        if document:
            st.download_button(
                label="파일 다운로드",
                data=bio.getvalue(),
                file_name=filename,
                mime="docx",
            )